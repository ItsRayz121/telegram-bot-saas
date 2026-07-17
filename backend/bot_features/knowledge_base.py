import math
import logging

logger = logging.getLogger(__name__)

PROVIDER_DEFAULT_MODELS = {
    "openai": "gpt-4o-mini",
    "openrouter": "openai/gpt-4o-mini",
    "custom": "gpt-4o-mini",
    "anthropic": "claude-haiku-4-5-20251001",
    "gemini": "gemini-1.5-flash",
}

PROVIDER_DEFAULT_BASE_URLS = {
    "openai": None,
    "openrouter": "https://openrouter.ai/api/v1",
    "anthropic": "https://api.anthropic.com",
    "gemini": "https://generativelanguage.googleapis.com/v1beta",
}


# The model is instructed (ai_personality._KNOWLEDGE_RULES) to output exactly this
# token when the KB context does not contain the answer. We convert it to
# answer=None so callers escalate to admins instead of posting "I don't know"
# into the group.
NO_ANSWER_SENTINEL = "NO_ANSWER"


def _is_no_answer(text):
    """True when the model signalled it has no answer (or produced nothing)."""
    if not text or not text.strip():
        return True
    stripped = text.strip()
    head = stripped.strip("\"'`*_ .!").upper()
    # Entire reply is the sentinel (possibly quoted/punctuated)…
    if head in (NO_ANSWER_SENTINEL, "NO ANSWER"):
        return True
    # …or the sentinel token followed by extra prose. Underscore form only, so a
    # real answer like "No answers are given for…" is never misclassified.
    return stripped.upper().startswith(NO_ANSWER_SENTINEL)


def _cosine_similarity(a, b):
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x * x for x in a))
    norm_b = math.sqrt(sum(x * x for x in b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


def _chunk_text(text, chunk_size=400, overlap=50):
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunks.append(" ".join(words[i:i + chunk_size]))
        i += chunk_size - overlap
    return [c for c in chunks if c.strip()]


def _extract_text(content_bytes, file_type):
    try:
        if file_type in ("txt", "md"):
            return content_bytes.decode("utf-8", errors="ignore")
        elif file_type == "pdf":
            import PyPDF2, io
            reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        elif file_type == "docx":
            import docx, io
            doc = docx.Document(io.BytesIO(content_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        logger.error(f"Text extraction error ({file_type}): {e}")
    return ""


class KnowledgeBaseSystem:

    def __init__(self, app):
        self.app = app

    def _load_group_api_key(self, group_id, telegram_group_id=None):
        """Resolve AI key using the central resolver. Returns dict or None."""
        try:
            with self.app.app_context():
                from ..assistant.ai_key_resolver import resolve_ai_provider_for_group, QuotaExceededError
                from ..models import Group, TelegramGroup

                # Resolve user_id from group
                user_id = None
                if telegram_group_id:
                    tg = TelegramGroup.query.filter_by(telegram_group_id=str(telegram_group_id)).first()
                    if tg:
                        user_id = tg.owner_user_id
                elif group_id:
                    grp = Group.query.get(group_id)
                    if grp:
                        from ..models import Bot
                        bot = Bot.query.get(grp.bot_id)
                        if bot:
                            user_id = bot.user_id

                if not user_id:
                    return None

                try:
                    result = resolve_ai_provider_for_group(user_id, group_id, telegram_group_id)
                except QuotaExceededError:
                    return None

                if not result.get("api_key"):
                    return None

                return {
                    "provider": result["provider"],
                    "api_key": result["api_key"],
                    "base_url": result.get("base_url"),
                    "model_name": result.get("model"),
                }
        except Exception as e:
            logger.error(f"Failed to load group API key: {e}")
            return None

    def _get_openai_client(self, api_key=None, base_url=None):
        """Return an OpenAI-compatible client."""
        from openai import OpenAI
        key = api_key
        if not key:
            return None
        kwargs = {"api_key": key}
        if base_url:
            kwargs["base_url"] = base_url
        return OpenAI(**kwargs)

    def _embed(self, texts, group_id=None, telegram_group_id=None):
        """Embed texts using group's custom key if available, else env key.

        OpenRouter does not support the embeddings API, so when the resolved
        key is an OpenRouter platform key we fall back to OPENAI_API_KEY.
        Users with an openai or custom key are used directly.
        """
        from .. import config as _cfg
        key_config = self._load_group_api_key(group_id, telegram_group_id) if (group_id or telegram_group_id) else None

        client = None
        if key_config and key_config["provider"] == "openai":
            client = self._get_openai_client(api_key=key_config["api_key"])
        elif key_config and key_config["provider"] == "custom":
            client = self._get_openai_client(
                api_key=key_config["api_key"],
                base_url=key_config.get("base_url"),
            )
        # OpenRouter cannot handle embeddings — fall through to env key
        if not client:
            client = self._get_openai_client(api_key=_cfg.Config.OPENAI_API_KEY or None)

        if not client:
            return [None] * len(texts)

        try:
            results = []
            for i in range(0, len(texts), 100):
                batch = texts[i:i + 100]
                resp = client.embeddings.create(model="text-embedding-3-small", input=batch)
                results.extend(e.embedding for e in resp.data)
            return results
        except Exception as e:
            logger.error(f"Embedding error: {e}")
            return [None] * len(texts)

    def process_document(self, group_id, filename, file_type, content_bytes, telegram_group_id=None):
        text = _extract_text(content_bytes, file_type)
        if not text.strip():
            return None, "Could not extract text from file"

        chunks = _chunk_text(text)
        if not chunks:
            return None, "File appears to be empty"

        embeddings = self._embed(chunks, group_id=group_id, telegram_group_id=telegram_group_id)
        if all(e is None for e in embeddings):
            return None, (
                "Embeddings failed: no AI API key is configured. "
                "Add your OpenAI API key in the Knowledge Base → AI Provider & API Key section, "
                "or set OPENAI_API_KEY on the server."
            )
        chunk_data = [
            {"text": c, "embedding": e}
            for c, e in zip(chunks, embeddings)
            if e is not None
        ]

        with self.app.app_context():
            from ..models import KnowledgeDocument, db
            doc = KnowledgeDocument(
                group_id=group_id,
                telegram_group_id=str(telegram_group_id) if telegram_group_id else None,
                filename=filename,
                file_type=file_type,
                content_text=text[:10000],
                chunks=chunk_data,
            )
            db.session.add(doc)
            db.session.commit()
            return doc.to_dict(), None

    def _log_kb_activity(self, question, group_id, telegram_group_id,
                         answered=True, confidence=None, source="knowledge_base",
                         answer=None):
        """Best-effort AI Activity log for a knowledge lookup. Never raises.

        Stores the generated answer in meta so the AI Activity tab can show a
        full question→answer preview (and the admin can correct it).
        """
        try:
            from ..ai_activity import log_ai_activity, derive_scope_ref
            scope, ref = derive_scope_ref(
                telegram_group_id=telegram_group_id, group_id=group_id
            )
            meta = {}
            if confidence is not None:
                meta["confidence"] = round(confidence, 3)
            if answer:
                meta["answer"] = str(answer)[:1000]
            with self.app.app_context():
                log_ai_activity(
                    scope, ref, "knowledge",
                    "FAQ answer generated" if answered else "Knowledge lookup (no answer)",
                    detail=(question or "")[:300],
                    status="ok" if answered else "skipped",
                    source=source, meta=(meta or None),
                )
        except Exception:
            pass

    async def answer_question(self, question, group_id, telegram_group_id=None,
                              group_name="this community", kb_settings=None,
                              auto_reply_triggers=None):
        """Returns (answer: str|None, confidence: float).

        auto_reply_triggers: optional list of {"trigger": str, "response": str} dicts
        injected as extra knowledge context when the admin has enabled that setting.
        """
        try:
            key_config = self._load_group_api_key(group_id, telegram_group_id)
            provider = key_config["provider"] if key_config else "openai"

            # Build trigger context block (safe, admin-authored content only)
            trigger_context = ""
            if auto_reply_triggers:
                snippets = [
                    f"Q: {t['trigger']}\nA: {t['response']}"
                    for t in auto_reply_triggers
                    if t.get("trigger") and t.get("response")
                ]
                if snippets:
                    trigger_context = "Known community Q&A:\n" + "\n\n".join(snippets)

            with self.app.app_context():
                from ..models import KnowledgeDocument
                if telegram_group_id:
                    docs = KnowledgeDocument.query.filter_by(
                        telegram_group_id=str(telegram_group_id)
                    ).all()
                else:
                    docs = KnowledgeDocument.query.filter_by(group_id=group_id).all()
                all_chunks = []
                for doc in docs:
                    for ch in (doc.chunks or []):
                        if ch.get("embedding"):
                            all_chunks.append(ch)

            # If no KB chunks but triggers exist, answer from triggers alone
            if not all_chunks and not trigger_context:
                logger.debug(f"KB: No chunks found for group {group_id or telegram_group_id}")
                return None, 0.0

            if not all_chunks and trigger_context:
                logger.debug(f"KB: No document chunks, using trigger context only for group {group_id or telegram_group_id}")
                key_config = self._load_group_api_key(group_id, telegram_group_id)
                if not key_config:
                    from .. import config as _cfg
                    fallback_key = _cfg.Config.OPENAI_API_KEY
                    if not fallback_key:
                        return None, 0.0
                answer = await self._generate_answer(
                    question, trigger_context, key_config,
                    group_name=group_name,
                    kb_settings=kb_settings,
                )
                if _is_no_answer(answer):
                    self._log_kb_activity(question, group_id, telegram_group_id,
                                          answered=False, source="auto_reply_triggers")
                    return None, 0.0
                self._log_kb_activity(question, group_id, telegram_group_id,
                                      answered=True, source="auto_reply_triggers",
                                      answer=answer)
                return answer, 0.75  # fixed confidence: admin-curated content

            # Embed question — OpenRouter cannot handle embeddings, use OpenAI key
            from .. import config as _cfg
            embed_client = None
            if key_config and provider == "openai":
                embed_client = self._get_openai_client(api_key=key_config["api_key"])
            elif key_config and provider == "custom":
                embed_client = self._get_openai_client(
                    api_key=key_config["api_key"],
                    base_url=key_config.get("base_url"),
                )
            if not embed_client:
                embed_client = self._get_openai_client(api_key=_cfg.Config.OPENAI_API_KEY or None)

            if not embed_client:
                logger.warning(
                    "KB: No AI API key configured for group %s — knowledge base Q&A unavailable. "
                    "Configure an API key in group AI Settings.",
                    group_id or telegram_group_id,
                )
                return None, 0.0

            q_resp = embed_client.embeddings.create(model="text-embedding-3-small", input=question)
            q_emb = q_resp.data[0].embedding

            scored = sorted(
                all_chunks,
                key=lambda c: _cosine_similarity(q_emb, c["embedding"]),
                reverse=True,
            )
            top_score = _cosine_similarity(q_emb, scored[0]["embedding"]) if scored else 0.0
            logger.debug(f"KB: Top confidence score for group {group_id}: {top_score:.3f}")

            # Retrieval floor: skip the LLM call only when nothing in the KB is
            # even vaguely related (saves tokens on chit-chat). The configured
            # confidence_threshold is capped at 0.25 because cosine similarity
            # with text-embedding-3-small rarely exceeds ~0.6 even for exact
            # matches — the old behaviour of gating the FINAL answer on this
            # score rejected correct answers. The model itself now decides
            # answerability via the NO_ANSWER sentinel.
            raw_threshold = (kb_settings or {}).get("confidence_threshold", 0.25)
            try:
                configured = float(raw_threshold)
            except (TypeError, ValueError):
                configured = 0.25
            floor = min(configured, 0.25)
            if top_score < floor and not trigger_context:
                self._log_kb_activity(question, group_id, telegram_group_id,
                                      answered=False, confidence=top_score,
                                      source="knowledge_base")
                return None, top_score

            context = "\n\n---\n\n".join(c["text"] for c in scored[:3])
            if trigger_context:
                context = context + "\n\n---\n\n" + trigger_context

            answer = await self._generate_answer(
                question, context, key_config,
                group_name=group_name,
                kb_settings=kb_settings,
            )
            if _is_no_answer(answer):
                self._log_kb_activity(question, group_id, telegram_group_id,
                                      answered=False, confidence=top_score,
                                      source="knowledge_base")
                return None, top_score
            self._log_kb_activity(question, group_id, telegram_group_id,
                                  answered=True, confidence=top_score,
                                  source="knowledge_base", answer=answer)
            return answer, top_score

        except Exception as e:
            logger.error(f"KB Q&A error: {e}")
            return None, 0.0

    async def _generate_answer(self, question, context, key_config,
                               group_name="this community", kb_settings=None):
        """Generate an answer from context using the appropriate provider."""
        from .ai_personality import build_system_prompt
        provider = key_config["provider"] if key_config else "openai"
        kb = kb_settings or {}

        system_prompt = build_system_prompt(
            personality_id=kb.get("personality", "professional_support"),
            group_name=group_name,
            custom_instructions=kb.get("custom_instructions", ""),
            reply_length=kb.get("reply_length", "balanced"),
            emoji_level=kb.get("emoji_level", "minimal"),
            formality_level=kb.get("formality_level", "neutral"),
        )
        user_prompt = f"Knowledge base context:\n{context}\n\nQuestion: {question}"

        if provider == "anthropic" and key_config:
            return await self._generate_anthropic(key_config, system_prompt, user_prompt)
        elif provider == "gemini" and key_config:
            return await self._generate_gemini(key_config, system_prompt, user_prompt)
        else:
            return await self._generate_openai_compatible(key_config, system_prompt, user_prompt, provider)

    async def _generate_openai_compatible(self, key_config, system_prompt, user_prompt, provider):
        from ..config import Config
        api_key = key_config["api_key"] if key_config else Config.OPENAI_API_KEY
        base_url = None
        model = "gpt-4o-mini"
        if key_config:
            base_url = key_config.get("base_url") or PROVIDER_DEFAULT_BASE_URLS.get(provider)
            model = key_config.get("model_name") or PROVIDER_DEFAULT_MODELS.get(provider, "gpt-4o-mini")

        client = self._get_openai_client(api_key=api_key, base_url=base_url)
        if not client:
            return None
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                max_tokens=400,
            )
            return resp.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"OpenAI-compatible generation error: {e}")
            return None

    async def _generate_anthropic(self, key_config, system_prompt, user_prompt):
        import requests as req
        model = key_config.get("model_name") or PROVIDER_DEFAULT_MODELS["anthropic"]
        base_url = key_config.get("base_url") or "https://api.anthropic.com"
        url = base_url.rstrip("/") + "/v1/messages"
        headers = {
            "x-api-key": key_config["api_key"],
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        }
        payload = {
            "model": model,
            "max_tokens": 400,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_prompt}],
        }
        try:
            resp = req.post(url, json=payload, headers=headers, timeout=30)
            if resp.status_code == 200:
                data = resp.json()
                return data.get("content", [{}])[0].get("text", "").strip()
        except Exception as e:
            logger.error(f"Anthropic generation error: {e}")
        return None

    async def _generate_gemini(self, key_config, system_prompt, user_prompt):
        import requests as req
        model = key_config.get("model_name") or PROVIDER_DEFAULT_MODELS["gemini"]
        api_key = key_config["api_key"]
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
        payload = {
            "contents": [{"parts": [{"text": f"{system_prompt}\n\n{user_prompt}"}]}],
            "generationConfig": {"maxOutputTokens": 400},
        }
        try:
            resp = req.post(url, json=payload, timeout=30)
            if resp.status_code == 200:
                data = resp.json()
                return data.get("candidates", [{}])[0].get("content", {}).get(
                    "parts", [{}])[0].get("text", "").strip()
        except Exception as e:
            logger.error(f"Gemini generation error: {e}")
        return None
