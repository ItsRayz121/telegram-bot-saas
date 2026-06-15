"""Knowledge-base document ingestion + semantic retrieval for Guildizer.

Parity with Telegizer's KnowledgeBaseSystem, adapted to Guildizer's stack:
  * extract text from uploaded PDF / DOCX / TXT / MD files,
  * split into overlapping word chunks,
  * embed each chunk with OpenAI ``text-embedding-3-small`` (platform key),
  * persist the chunks (+ embeddings) on the KnowledgeDocument row,
  * answer /ask via cosine similarity over those chunks.

Guildizer is intentionally a *global-key* product (one platform OPENAI_API_KEY
on the worker, no per-guild keys), so embeddings always use Config.OPENAI_API_KEY.
When no key is set, documents are still stored and indexed for the keyword
fallback in ``knowledge.py`` — uploads never hard-fail just because vision/AI is
off. Self-contained: no Telegizer imports (Discord isolation rule).
"""
from __future__ import annotations

import io
import logging
import math
import re

from config import Config
from database import SessionLocal
from models import KnowledgeDocument

log = logging.getLogger("guildizer.kb")

EMBED_MODEL = "text-embedding-3-small"
ALLOWED_EXT = ("pdf", "docx", "txt", "md")
MAX_FILE_BYTES = 5 * 1024 * 1024            # 5 MB per file (matches Telegizer)
_MAX_CONTENT_TEXT = 10000                   # stored extracted-text cap
_TOP_K = 3
_MAX_CONTEXT = 4000

# Magic-byte signatures — extension alone is spoofable.
_MAGIC_SIGNATURES = {
    b"%PDF": "pdf",
    b"PK\x03\x04": "docx",   # docx is a ZIP container
}


# --- extraction --------------------------------------------------------------

def _extract_text(content_bytes: bytes, file_type: str) -> str:
    try:
        if file_type in ("txt", "md"):
            return content_bytes.decode("utf-8", errors="ignore")
        if file_type == "pdf":
            from pypdf import PdfReader  # lazy: optional dependency
            reader = PdfReader(io.BytesIO(content_bytes))
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        if file_type == "docx":
            import docx  # lazy: python-docx
            d = docx.Document(io.BytesIO(content_bytes))
            return "\n".join(p.text for p in d.paragraphs)
    except Exception as e:  # noqa: BLE001
        log.error("KB text extraction failed (%s): %s", file_type, e)
    return ""


def _chunk_text(text: str, chunk_size: int = 400, overlap: int = 50) -> list[str]:
    words = (text or "").split()
    chunks, i = [], 0
    while i < len(words):
        chunks.append(" ".join(words[i:i + chunk_size]))
        i += chunk_size - overlap
    return [c for c in chunks if c.strip()]


# --- embeddings --------------------------------------------------------------

def _embed(texts: list[str]) -> list[list[float] | None]:
    """Embed texts with the platform OpenAI key. Returns one vector per text, or
    None per text when embeddings are unavailable (no key / SDK / API error).
    Never raises — callers degrade to keyword search."""
    if not texts:
        return []
    if not Config.OPENAI_API_KEY:
        return [None] * len(texts)
    try:
        from openai import OpenAI  # lazy
    except ImportError:
        log.warning("openai SDK not installed; KB embeddings disabled")
        return [None] * len(texts)
    client = OpenAI(api_key=Config.OPENAI_API_KEY)
    out: list[list[float] | None] = []
    # Batch to stay well under request limits.
    for start in range(0, len(texts), 64):
        batch = texts[start:start + 64]
        try:
            resp = client.embeddings.create(model=EMBED_MODEL, input=batch)
            out.extend(e.embedding for e in resp.data)
        except Exception as e:  # noqa: BLE001
            log.error("KB embedding error: %s", e)
            out.extend([None] * len(batch))
    return out


def _cosine(a: list[float], b: list[float]) -> float:
    if not a or not b:
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(y * y for y in b))
    return dot / (na * nb) if na and nb else 0.0


# --- ingestion ---------------------------------------------------------------

def detect_type(filename: str, content_bytes: bytes) -> tuple[str | None, str | None]:
    """(file_type, error). Validates extension against magic bytes."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    if ext not in ALLOWED_EXT:
        return None, f"Unsupported file type .{ext}. Use: {', '.join(ALLOWED_EXT)}"
    detected = None
    for magic, ftype in _MAGIC_SIGNATURES.items():
        if content_bytes[:len(magic)] == magic:
            detected = ftype
            break
    if ext in ("txt", "md"):
        try:
            content_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return None, "File does not appear to be valid UTF-8 text"
    elif detected is None:
        return None, f"File does not appear to be a valid {ext.upper()}"
    elif detected != ext:
        return None, f"File content does not match its extension (.{ext})"
    return ext, None


def process_upload(db, guild_id: int, filename: str, file_type: str,
                   content_bytes: bytes) -> tuple[dict | None, str | None]:
    """Extract → chunk → embed → persist. Returns (doc_dict, error)."""
    text = _extract_text(content_bytes, file_type)
    if not text.strip():
        return None, "Could not extract any text from this file."
    chunks = _chunk_text(text)
    if not chunks:
        return None, "No usable text after processing."
    embeddings = _embed(chunks)
    chunk_data = [{"text": c, "embedding": e} for c, e in zip(chunks, embeddings)]

    row = KnowledgeDocument(
        guild_id=guild_id,
        title=filename[:200],
        file_type=file_type,
        content="",                       # manual-text field unused for uploads
        content_text=text[:_MAX_CONTENT_TEXT],
        chunks=chunk_data,
        enabled=True,
    )
    db.add(row)
    db.commit()
    return row.to_dict(), None


# --- retrieval ---------------------------------------------------------------

def semantic_context(guild_id: int, question: str) -> tuple[str | None, bool] | None:
    """Best-effort semantic retrieval over embedded chunks.

    Returns ``(context, confident)`` when embedded chunks exist, or ``None`` to
    signal the caller should use the keyword fallback (no embeddings available)."""
    db = SessionLocal()
    try:
        docs = (
            db.query(KnowledgeDocument)
            .filter(KnowledgeDocument.guild_id == guild_id,
                    KnowledgeDocument.enabled.is_(True))
            .limit(50)
            .all()
        )
        all_chunks = []
        for d in docs:
            for ch in (d.chunks or []):
                if ch.get("embedding"):
                    all_chunks.append(ch)
        if not all_chunks:
            return None  # no embedded chunks → let keyword search handle it
        q_emb = _embed([question])[0]
        if q_emb is None:
            return None
        scored = sorted(
            ((_cosine(q_emb, ch["embedding"]), ch) for ch in all_chunks),
            key=lambda x: x[0], reverse=True,
        )
        confident = scored[0][0] >= 0.30
        context = ""
        for _score, ch in scored[:_TOP_K]:
            piece = ch["text"] + "\n\n"
            if len(context) + len(piece) > _MAX_CONTEXT:
                break
            context += piece
        return context.strip(), confident
    finally:
        db.close()
        SessionLocal.remove()
